"""
Tạo file đề bài và rubric chấm điểm cho môn Kĩ năng dạy học
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "sample_assignments")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    if align:
        p.alignment = align
    return p


# ==========================================================
# FILE 1: ĐỀ BÀI
# ==========================================================
def create_assignment_prompt():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Header quốc huy
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h1.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC SƯ PHẠM TP. HỒ CHÍ MINH\nKHOA GIÁO DỤC TIỂU HỌC")
    r.bold = True; r.font.size = Pt(13)

    doc.add_paragraph()

    # Tiêu đề
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("ĐỀ BÀI KIỂM TRA GIỮA KỲ")
    t.bold = True; t.font.size = Pt(16)
    t.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Môn: KỸ NĂNG DẠY HỌC")
    s.bold = True; s.font.size = Pt(14)
    s.font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    doc.add_paragraph()

    # Thông tin môn học
    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = 'Table Grid'
    data = [
        ("Học kỳ:", "I - Năm học 2025-2026", "Hình thức:", "Bài tập cá nhân"),
        ("Ngày giao:", "01/06/2026", "Ngày nộp:", "03/06/2026"),
        ("Lớp:", "GDTH-K21 (A, B, C)", "Số TC:", "3 tín chỉ"),
        ("GV phụ trách:", "TS. Trần Văn Bình", "Thang điểm:", "10 điểm"),
        ("Hình thức nộp:", "File DOCX + bản in, nộp tại văn phòng khoa", "", ""),
    ]
    for r_idx, row_data in enumerate(data):
        row = info_table.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            row.cells[c_idx].text = text
            if c_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], "D6EAF8")
                for para in row.cells[c_idx].paragraphs:
                    for run in para.runs:
                        run.bold = True; run.font.size = Pt(10)
            else:
                for para in row.cells[c_idx].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
    doc.add_paragraph()

    # PHẦN ĐƯỜNG KẺ PHÂN CÁCH
    line = doc.add_paragraph("━" * 65)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # NỘI DUNG ĐỀ BÀI
    h = doc.add_heading("NỘI DUNG BÀI TẬP", 1)
    h.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    add_para(doc, "THIẾT KẾ GIÁO ÁN DẠY HỌC MÔN TOÁN LỚP 5", bold=True, size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Chủ đề: HÌNH HỌC VÀ ĐO LƯỜNG", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Yêu cầu
    req = doc.add_heading("I. YÊU CẦU NHIỆM VỤ", 2)
    req.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    add_para(doc, "Sinh viên thiết kế 01 giáo án dạy học hoàn chỉnh cho 01 bài học "
             "thuộc mạch nội dung Hình học và Đo lường, môn Toán lớp 5, theo Chương trình "
             "GDPT 2018. Giáo án phải đảm bảo các yêu cầu sau:", size=11)

    reqs = [
        "Chọn 1 bài học cụ thể trong Chương trình Toán lớp 5 thuộc mạch Hình học (ví dụ: Diện tích hình thang, Diện tích hình tròn, Thể tích hình hộp chữ nhật, Diện tích xung quanh và diện tích toàn phần,…)",
        "Xây dựng mục tiêu bài học rõ ràng theo 3 mặt: Kiến thức, Kỹ năng, Thái độ (hoặc theo Năng lực, Phẩm chất theo CTGDPT 2018)",
        "Thiết kế ít nhất 4 hoạt động dạy học trong tiến trình lên lớp, có phân bổ thời gian hợp lý",
        "Mỗi hoạt động cần nêu rõ: Mục tiêu hoạt động, Nội dung, Phương pháp/kỹ thuật dạy học, Đồ dùng/phương tiện, Cách đánh giá",
        "Đề xuất ít nhất 2 phương pháp dạy học tích cực và giải thích lý do chọn",
        "Có kế hoạch đánh giá học sinh rõ ràng (cả quá trình và kết quả)",
        "Dự kiến ít nhất 2 tình huống sư phạm có thể xảy ra và cách xử lý",
        "Có danh mục tài liệu tham khảo (tối thiểu 3 nguồn)",
    ]
    for i, r in enumerate(reqs, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(r).font.size = Pt(11)

    doc.add_paragraph()

    # CẤU TRÚC GIÁO ÁN
    struct = doc.add_heading("II. CẤU TRÚC GIÁO ÁN ĐƯỢC GỢI Ý", 2)
    struct.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    struct_table = doc.add_table(rows=1, cols=3)
    struct_table.style = 'Table Grid'
    for i, h_text in enumerate(["Phần", "Nội dung", "Lưu ý"]):
        struct_table.rows[0].cells[i].text = h_text
        set_cell_bg(struct_table.rows[0].cells[i], "1A5F7A")
        for para in struct_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    struct_data = [
        ("Phần I:\nThông tin chung", "Tên bài, lớp, thời gian, SGK sử dụng, ngày dạy, giáo viên dạy", "Đủ, chính xác"),
        ("Phần II:\nMục tiêu", "Kiến thức, Kỹ năng, Thái độ/Năng lực, Phẩm chất. Dùng động từ đo lường được.", "Dùng thang Bloom"),
        ("Phần III:\nChuẩn bị", "Đồ dùng GV, đồ dùng HS, phương tiện kỹ thuật", "Thực tế, khả thi"),
        ("Phần IV:\nTiến trình", "Các hoạt động dạy học theo thứ tự: Khởi động → Khám phá → Thực hành → Vận dụng → Tổng kết", "Chi tiết, rõ ràng"),
        ("Phần V:\nTình huống SP", "Tình huống dự kiến + Cách xử lý", "Thực tế, sáng tạo"),
        ("Phần VI:\nĐánh giá", "Đánh giá quá trình và kết quả, tiêu chí đánh giá", "Đa dạng, cụ thể"),
        ("Phần VII:\nTài liệu", "Danh mục tham khảo theo chuẩn APA hoặc chuẩn Việt Nam", "≥ 3 tài liệu"),
    ]
    row_colors = ["FDFEFE", "EBF5FB", "FDFEFE", "EBF5FB", "FDFEFE", "EBF5FB", "FDFEFE"]
    for k, (part, content, note) in enumerate(struct_data):
        row = struct_table.add_row()
        row.cells[0].text = part; row.cells[1].text = content; row.cells[2].text = note
        set_cell_bg(row.cells[0], row_colors[k])
        for j, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if j == 0:
                        run.bold = True
    doc.add_paragraph()

    # QUY ĐỊNH
    rule = doc.add_heading("III. QUY ĐỊNH TRÌNH BÀY", 2)
    rule.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)
    
    rules = [
        "Font chữ: Times New Roman, cỡ 13, giãn dòng 1.5",
        "Lề: Trên/Dưới: 2cm, Trái: 3cm, Phải: 2cm",
        "Giáo án không ít hơn 5 trang A4 (không tính trang bìa)",
        "Có trang bìa đầy đủ thông tin: Tên trường, Khoa, Tên sinh viên, MSSV, Lớp, Tên môn học, GV hướng dẫn",
        "File nộp đặt tên: HoTen_MSSV_GiaoAn.docx (ví dụ: NguyenThiAn_21CDTH001_GiaoAn.docx)",
        "Bài làm phải là công trình cá nhân, không sao chép. Phát hiện đạo văn sẽ không được chấm điểm.",
    ]
    for r in rules:
        p = doc.add_paragraph(f"▪ {r}")
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # THANG ĐIỂM
    grade = doc.add_heading("IV. THANG ĐIỂM TỔNG QUÁT", 2)
    grade.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    grade_table = doc.add_table(rows=1, cols=5)
    grade_table.style = 'Table Grid'
    for i, h_text in enumerate(["Tiêu chí", "Trọng số", "Xuất sắc\n(9-10)", "Đạt\n(5-8)", "Không đạt\n(< 5)"]):
        grade_table.rows[0].cells[i].text = h_text
        set_cell_bg(grade_table.rows[0].cells[i], "C0392B")
        for para in grade_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    
    grade_data = [
        ("Mục tiêu bài học", "20%", "Đầy đủ 3 mặt, dùng động từ cụ thể, đo lường được", "Có ≥ 2 mặt, tương đối rõ", "Sơ sài, chung chung"),
        ("Tiến trình dạy học", "40%", "≥ 4 HĐ, đủ chi tiết, PPDH tích cực, phân bổ TG hợp lý", "3-4 HĐ, tương đối chi tiết", "< 3 HĐ, thiếu chi tiết"),
        ("Phương pháp DH", "15%", "≥ 2 PP tích cực, giải thích thuyết phục", "1-2 PP, giải thích đơn giản", "Chỉ dùng PP truyền thống"),
        ("Đánh giá HS", "15%", "Đa dạng, có tiêu chí rõ ràng, cả QT và KQ", "Có đánh giá nhưng chưa đa dạng", "Thiếu hoặc không có"),
        ("Trình bày & Tài liệu", "10%", "Đúng quy định, có ≥ 3 TLTK đúng chuẩn", "Tương đối đúng quy định", "Sai quy định, thiếu TLTK"),
    ]
    row_colors2 = ["D5F5E3", "D6EAF8", "FEF9E7", "F9EBEA", "EAF2FF"]
    for k, (crit, weight, ex, ok, fail) in enumerate(grade_data):
        row = grade_table.add_row()
        row.cells[0].text = crit; row.cells[1].text = weight
        row.cells[2].text = ex; row.cells[3].text = ok; row.cells[4].text = fail
        set_cell_bg(row.cells[0], row_colors2[k])
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_para.add_run("⚠️ Lưu ý: ").bold = True
    note_para.add_run("Rubric chấm điểm chi tiết được cung cấp kèm theo tài liệu riêng. "
                      "Sinh viên cần đọc kỹ rubric trước khi làm bài.")
    doc.add_paragraph()

    # Ký tên GV
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign.add_run("TP.HCM, ngày 01 tháng 06 năm 2026\nGIẢNG VIÊN PHỤ TRÁCH MÔN\n\n\n\nTS. Trần Văn Bình")
    sign_run.bold = True; sign_run.font.size = Pt(12)

    fname = os.path.join(OUTPUT_DIR, "00_DE_BAI_KyNangDayHoc.docx")
    doc.save(fname)
    print(f"✅ Đã tạo đề bài: {fname}")


# ==========================================================
# FILE 2: RUBRIC CHẤM ĐIỂM
# ==========================================================
def create_rubric():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.page_width = Cm(29.7)  # A4 landscape
    section.page_height = Cm(21)

    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h1.add_run("TRƯỜNG ĐẠI HỌC SƯ PHẠM TP. HỒ CHÍ MINH - KHOA GIÁO DỤC TIỂU HỌC\n")
    r.bold = True; r.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("BẢNG RUBRIC CHẤM ĐIỂM GIÁO ÁN DẠY TOÁN LỚP 5")
    t.bold = True; t.font.size = Pt(15)
    t.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Môn: Kĩ năng dạy học | Trọng số: Điểm giữa kỳ | Tổng điểm tối đa: 10.0 điểm").font.size = Pt(11)
    doc.add_paragraph()

    # RUBRIC TABLE
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    headers = [
        "STT", "TIÊU CHÍ ĐÁNH GIÁ", "TRỌNG SỐ",
        "MỨC 4 - XUẤT SẮC\n(90-100% điểm TC)",
        "MỨC 3 - TỐT\n(70-89% điểm TC)",
        "MỨC 2 - ĐẠT\n(50-69% điểm TC)",
        "MỨC 1 - CHƯA ĐẠT\n(< 50% điểm TC)"
    ]
    header_row = table.rows[0]
    for i, h_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h_text
        set_cell_bg(cell, "1A3A5C")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    rubric_data = [
        # (STT, Tiêu chí, Trọng số, Mức 4, Mức 3, Mức 2, Mức 1)
        ("1", "MỤC TIÊU BÀI HỌC\n(Kiến thức, Kỹ năng, Thái độ/NL)",
         "20%\n(2.0đ)",
         "• Trình bày đầy đủ cả 3 mặt (KT-KN-TĐ)\n• Dùng động từ hành động đo lường được\n• Mục tiêu cụ thể, phù hợp trình độ HS lớp 5\n• Có mục tiêu về NL và PC theo CT 2018\n→ 1.8 – 2.0 điểm",
         "• Đủ 3 mặt nhưng một mặt chưa cụ thể\n• Hầu hết động từ đo lường được\n• Phù hợp HS lớp 5\n→ 1.4 – 1.7 điểm",
         "• Chỉ có 1-2 mặt\n• Một số động từ còn chung chung\n• Chưa phân biệt rõ các mặt\n→ 1.0 – 1.3 điểm",
         "• Thiếu mục tiêu hoặc quá chung\n• Không dùng động từ đo lường\n• Không phù hợp trình độ\n→ < 1.0 điểm"),

        ("2", "CHUẨN BỊ DẠY - HỌC\n(Đồ dùng, phương tiện)",
         "5%\n(0.5đ)",
         "• Liệt kê đầy đủ, chi tiết đồ dùng GV và HS\n• Đồ dùng sáng tạo, hỗ trợ tốt mục tiêu\n• Có phương tiện kỹ thuật phù hợp\n→ 0.45 – 0.5 điểm",
         "• Đủ đồ dùng GV và HS\n• Phù hợp với bài học\n→ 0.35 – 0.44 điểm",
         "• Chỉ liệt kê đồ dùng cơ bản\n• Thiếu chi tiết\n→ 0.25 – 0.34 điểm",
         "• Liệt kê quá sơ sài (chỉ 'SGK, bảng')\n• Không phù hợp bài học\n→ < 0.25 điểm"),

        ("3", "TIẾN TRÌNH DẠY HỌC\n(Số lượng và chất lượng HĐ)",
         "20%\n(2.0đ)",
         "• ≥ 4 hoạt động rõ ràng theo chu trình\n• Mỗi HĐ đủ: mục tiêu, nội dung, PP, TG\n• Phân bổ thời gian hợp lý, logic\n• Có chuyển tiếp giữa các HĐ tự nhiên\n→ 1.8 – 2.0 điểm",
         "• 4 hoạt động, phần lớn đủ chi tiết\n• Phân bổ TG tương đối hợp lý\n→ 1.4 – 1.7 điểm",
         "• 3 hoạt động, thiếu chi tiết 1 số HĐ\n• TG chưa hợp lý ở một số HĐ\n→ 1.0 – 1.3 điểm",
         "• < 3 hoạt động\n• Thiếu chi tiết, không rõ ràng\n• Không có phân bổ TG\n→ < 1.0 điểm"),

        ("4", "PHƯƠNG PHÁP DẠY HỌC TÍCH CỰC\n(Loại PP và mức độ phù hợp)",
         "15%\n(1.5đ)",
         "• Sử dụng ≥ 2 PPDH tích cực (TLN, dự án, khám phá, PBL…)\n• Giải thích rõ lý do chọn PP\n• PP phù hợp với mục tiêu và đối tượng HS\n• Kỹ thuật DH cụ thể (Think-Pair-Share, KWL…)\n→ 1.35 – 1.5 điểm",
         "• 2 PPDH tích cực, giải thích sơ lược\n• Phù hợp với bài học\n→ 1.05 – 1.34 điểm",
         "• 1 PPDH tích cực + DH truyền thống\n• Chưa giải thích rõ\n→ 0.75 – 1.04 điểm",
         "• Chỉ dùng PP thuyết trình, hỏi đáp đơn thuần\n• Không có PPDH tích cực\n→ < 0.75 điểm"),

        ("5", "SỬ DỤNG CÂU HỎI\n(Chất lượng câu hỏi trong HĐ)",
         "10%\n(1.0đ)",
         "• Câu hỏi đa dạng: nhớ, hiểu, vận dụng, phân tích (Bloom)\n• Câu hỏi mở, kích thích tư duy\n• Phân bố đều trong các HĐ\n→ 0.9 – 1.0 điểm",
         "• Câu hỏi có ở các cấp độ nhớ-hiểu-VD\n• Một số câu hỏi mở\n→ 0.7 – 0.89 điểm",
         "• Chủ yếu câu hỏi đóng, cấp độ thấp\n• Ít câu hỏi kích thích tư duy\n→ 0.5 – 0.69 điểm",
         "• Hầu hết là câu hỏi đóng 'có/không'\n• Không có câu hỏi kích thích tư duy\n→ < 0.5 điểm"),

        ("6", "KẾ HOẠCH ĐÁNH GIÁ HS\n(Đánh giá QT và KQ)",
         "15%\n(1.5đ)",
         "• Đánh giá quá trình + kết quả đầy đủ\n• Đa dạng hình thức: vấn đáp, phiếu học tập, sản phẩm, tự đánh giá\n• Có tiêu chí đánh giá cụ thể\n• Phù hợp với mục tiêu đề ra\n→ 1.35 – 1.5 điểm",
         "• Có đánh giá QT và KQ\n• 2-3 hình thức đánh giá\n• Tiêu chí tương đối rõ\n→ 1.05 – 1.34 điểm",
         "• Chỉ đánh giá kết quả\n• 1-2 hình thức\n• Tiêu chí chung chung\n→ 0.75 – 1.04 điểm",
         "• Không có kế hoạch đánh giá\n• Hoặc chỉ ghi 'chấm bài'\n→ < 0.75 điểm"),

        ("7", "TÌNH HUỐNG SƯ PHẠM\n(Dự kiến và xử lý)",
         "5%\n(0.5đ)",
         "• ≥ 2 tình huống dự kiến\n• Tình huống thực tế, có khả năng xảy ra\n• Cách xử lý sư phạm, phù hợp tâm lý HS\n→ 0.45 – 0.5 điểm",
         "• 2 tình huống, xử lý tương đối hợp lý\n→ 0.35 – 0.44 điểm",
         "• 1 tình huống, xử lý đơn giản\n→ 0.25 – 0.34 điểm",
         "• Không có tình huống sư phạm\n→ < 0.25 điểm"),

        ("8", "TRÌNH BÀY VÀ TÀI LIỆU\n(Form + TLTK)",
         "10%\n(1.0đ)",
         "• Đúng quy định định dạng\n• Trang bìa đầy đủ, chuyên nghiệp\n• ≥ 3 TLTK đúng chuẩn APA hoặc Việt Nam\n• Không lỗi chính tả, ngữ pháp\n→ 0.9 – 1.0 điểm",
         "• Đúng quy định cơ bản\n• ≥ 3 TLTK\n• Ít lỗi chính tả\n→ 0.7 – 0.89 điểm",
         "• Một số lỗi định dạng\n• 2 TLTK hoặc chưa đúng chuẩn\n• Có lỗi chính tả\n→ 0.5 – 0.69 điểm",
         "• Sai nhiều quy định định dạng\n• < 2 TLTK hoặc không có\n• Nhiều lỗi chính tả\n→ < 0.5 điểm"),
    ]

    level_colors = ["1E8449", "2980B9", "F39C12", "C0392B"]
    row_alt = ["FDFEFE", "F0F3F4"]

    for k, row_data in enumerate(rubric_data):
        row = table.add_row()
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j == 0:
                set_cell_bg(cell, "2C3E50")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                        run.font.size = Pt(11)
            elif j == 1:
                set_cell_bg(cell, row_alt[k % 2])
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True; run.font.size = Pt(9)
            elif j == 2:
                set_cell_bg(cell, "FAD7A0")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True; run.font.size = Pt(10)
            else:
                # Mức 4, 3, 2, 1
                level_idx = j - 3
                bg = ["D5F5E3", "D6EAF8", "FEF9E7", "FDEDEC"][level_idx]
                set_cell_bg(cell, bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    doc.add_paragraph()

    # TỔNG KẾT ĐIỂM
    sum_para = doc.add_paragraph()
    sum_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sum_para.add_run("TỔNG ĐIỂM: 2.0 + 0.5 + 2.0 + 1.5 + 1.0 + 1.5 + 0.5 + 1.0 = 10.0 ĐIỂM")
    sr.bold = True; sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_paragraph()

    # BẢNG GHI ĐIỂM
    grade_h = doc.add_heading("BẢNG GHI ĐIỂM (Dành cho giảng viên)", 2)
    grade_h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    score_table = doc.add_table(rows=1, cols=6)
    score_table.style = 'Table Grid'
    for i, h_text in enumerate(["STT", "Họ và tên SV", "MSSV", "Điểm từng tiêu chí\n(TC1/TC2/.../TC8)", "Tổng điểm", "Ghi chú"]):
        score_table.rows[0].cells[i].text = h_text
        set_cell_bg(score_table.rows[0].cells[i], "2C3E50")
        for para in score_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
    for _ in range(15):
        row = score_table.add_row()
        for j in range(6):
            row.cells[j].text = ""
    doc.add_paragraph()

    # Ghi chú
    note = doc.add_paragraph()
    note.add_run("GHI CHÚ:\n").bold = True
    notes = [
        "Điểm làm tròn đến 0.25 điểm (ví dụ: 7.25, 7.5, 7.75)",
        "Bài nộp trễ trừ 10% điểm/ngày trễ, tối đa trừ 30%",
        "Đạo văn ≥ 40% (kiểm tra bằng phần mềm): Điểm = 0",
        "Tiêu chí TRỌNG YẾU: TC1 (Mục tiêu) + TC3 (Tiến trình) + TC4 (PPDH) — chiếm 55% điểm",
    ]
    for n in notes:
        note.add_run(f"• {n}\n").font.size = Pt(10)
    doc.add_paragraph()

    # Legend
    legend_h = doc.add_heading("CHÚ THÍCH MÀU SẮC", 2)
    legend_h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)
    
    leg_table = doc.add_table(rows=1, cols=4)
    leg_table.style = 'Table Grid'
    levels_info = [
        ("MỨC 4 - XUẤT SẮC", "D5F5E3", "Vượt chuẩn, thể hiện tư duy sư phạm sáng tạo, bài làm có chiều sâu"),
        ("MỨC 3 - TỐT", "D6EAF8", "Đạt chuẩn tốt, có thể hiện được hiểu biết về PP dạy học"),
        ("MỨC 2 - ĐẠT", "FEF9E7", "Đạt chuẩn tối thiểu, còn một số hạn chế cần cải thiện"),
        ("MỨC 1 - CHƯA ĐẠT", "FDEDEC", "Chưa đạt chuẩn, cần chỉnh sửa cơ bản trước khi nộp lại"),
    ]
    for i, (level, color, desc) in enumerate(levels_info):
        col_cells = leg_table.rows[0].cells
        col_cells[i].text = f"{level}\n{desc}"
        set_cell_bg(col_cells[i], color)
        for para in col_cells[i].paragraphs:
            for j_run, run in enumerate(para.runs):
                run.font.size = Pt(9)
                if j_run == 0:
                    run.bold = True

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run("TP.HCM, ngày 01 tháng 06 năm 2026\nGiảng viên phụ trách\n\n\n\nTS. Trần Văn Bình").bold = True

    fname = os.path.join(OUTPUT_DIR, "00_RUBRIC_ChamDiem_GiaoAn.docx")
    doc.save(fname)
    print(f"✅ Đã tạo rubric: {fname}")


if __name__ == "__main__":
    print("📋 Đang tạo đề bài và rubric...\n")
    create_assignment_prompt()
    create_rubric()
    print(f"\n✅ Hoàn thành! File lưu tại: {OUTPUT_DIR}")
