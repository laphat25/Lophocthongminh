"""
Tạo các file DOCX mẫu bài làm sinh viên cho môn Kĩ năng dạy học
Chủ đề: Thiết kế giáo án dạy Toán lớp 5
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "sample_assignments")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_bg(cell, hex_color):
    """Set background color for table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_row(table, cells_data, bg_colors=None, bold_row=False):
    row = table.add_row()
    for i, (cell_data) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = cell_data
        if bg_colors and i < len(bg_colors) and bg_colors[i]:
            set_cell_bg(cell, bg_colors[i])
        if bold_row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(10)
    return row


# ============================================================
# BÀI LÀM 1: Xuất sắc - Nguyễn Thị An (9.5 điểm)
# ============================================================
def create_assignment_excellent():
    doc = Document()
    
    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("TRƯỜNG ĐẠI HỌC SƯ PHẠM TP. HỒ CHÍ MINH\n")
    run.bold = True
    run.font.size = Pt(13)
    run2 = header_para.add_run("KHOA GIÁO DỤC TIỂU HỌC\n")
    run2.bold = True
    run2.font.size = Pt(12)

    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_para.add_run("GIÁO ÁN DẠY HỌC MÔN TOÁN LỚP 5")
    t.bold = True
    t.font.size = Pt(16)
    t.font.color.rgb = RGBColor(0x1a, 0x5f, 0x7a)

    doc.add_paragraph()

    # Student info table
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ["Họ và tên:", "Nguyễn Thị An", "MSSV:", "21CDTH001"],
        ["Lớp:", "GDTH-K21A", "Môn học:", "Kĩ năng dạy học"],
        ["Giảng viên:", "TS. Trần Văn Bình", "Ngày nộp:", "03/06/2026"],
        ["Bài:", "Thiết kế giáo án Toán lớp 5", "", ""],
    ]
    for r_idx, row_data in enumerate(info_data):
        row = info_table.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if c_idx % 2 == 0:
                        run.bold = True
            if c_idx % 2 == 0:
                set_cell_bg(cell, "D6EAF8")
    doc.add_paragraph()

    # PHẦN I: THÔNG TIN GIÁO ÁN
    add_heading(doc, "PHẦN I: THÔNG TIN CHUNG VỀ BÀI HỌC", 1, (0x1a, 0x5f, 0x7a))
    
    lesson_table = doc.add_table(rows=7, cols=2)
    lesson_table.style = 'Table Grid'
    lesson_table.columns[0].width = Cm(5)
    lesson_info = [
        ("Tên bài học:", "Diện tích hình thang (Toán lớp 5 - Tập 1, Bài 31)"),
        ("Lớp:", "5A - Trường Tiểu học Lê Lợi, Quận 3, TP.HCM"),
        ("Thời gian:", "45 phút (1 tiết học)"),
        ("Ngày dạy:", "Thứ Tư, 10/09/2026"),
        ("Phòng học:", "Lớp học thông thường (35 học sinh)"),
        ("Sách giáo khoa:", "SGK Toán 5 - Bộ Cánh Diều (2024)"),
        ("Bài trước:", "Diện tích hình thoi (Bài 30)"),
    ]
    for key, val in lesson_info:
        row = lesson_table.add_row()
        row.cells[0].text = key
        row.cells[1].text = val
        set_cell_bg(row.cells[0], "EBF5FB")
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

    # MỤC TIÊU
    add_heading(doc, "PHẦN II: MỤC TIÊU BÀI HỌC", 1, (0x1a, 0x5f, 0x7a))
    
    add_para(doc, "1. Về kiến thức:", bold=True)
    doc.add_paragraph("• Học sinh nhận biết được đặc điểm của hình thang (có 2 cạnh song song).", style='List Bullet')
    doc.add_paragraph("• Học sinh ghi nhớ và vận dụng được công thức tính diện tích hình thang: S = (a + b) × h ÷ 2.", style='List Bullet')
    doc.add_paragraph("• Học sinh phân biệt được các yếu tố: đáy lớn (a), đáy bé (b), chiều cao (h).", style='List Bullet')

    add_para(doc, "2. Về kỹ năng:", bold=True)
    doc.add_paragraph("• Tính toán chính xác diện tích hình thang trong bài tập đơn giản và bài toán có lời văn.", style='List Bullet')
    doc.add_paragraph("• Đo đạc và xác định các yếu tố của hình thang bằng thước.", style='List Bullet')
    doc.add_paragraph("• Vẽ hình thang và ghi chú các yếu tố cần thiết.", style='List Bullet')

    add_para(doc, "3. Về thái độ và phẩm chất:", bold=True)
    doc.add_paragraph("• Tự giác, tích cực tham gia các hoạt động học tập.", style='List Bullet')
    doc.add_paragraph("• Hợp tác hiệu quả trong nhóm, lắng nghe và tôn trọng ý kiến bạn.", style='List Bullet')
    doc.add_paragraph("• Yêu thích môn Toán, thấy được ứng dụng thực tế của kiến thức.", style='List Bullet')

    add_para(doc, "4. Năng lực hình thành:", bold=True)
    doc.add_paragraph("• Năng lực tư duy toán học: Phân tích, suy luận từ hình học đến công thức.", style='List Bullet')
    doc.add_paragraph("• Năng lực giải quyết vấn đề: Vận dụng công thức vào bài toán thực tế.", style='List Bullet')
    doc.add_paragraph("• Năng lực giao tiếp toán học: Diễn đạt cách tính, trình bày bài giải rõ ràng.", style='List Bullet')
    doc.add_paragraph()

    # CHUẨN BỊ
    add_heading(doc, "PHẦN III: CHUẨN BỊ DẠY - HỌC", 1, (0x1a, 0x5f, 0x7a))
    
    add_para(doc, "1. Giáo viên:", bold=True)
    doc.add_paragraph("• Giáo án chi tiết, bảng phụ có vẽ hình thang với kích thước rõ ràng.", style='List Bullet')
    doc.add_paragraph("• Mô hình hình thang bằng bìa cứng (3 loại kích thước khác nhau).", style='List Bullet')
    doc.add_paragraph("• Phiếu học tập nhóm (in sẵn, 7 phiếu cho 7 nhóm).", style='List Bullet')
    doc.add_paragraph("• Máy chiếu, file PowerPoint minh họa (12 slide).", style='List Bullet')
    doc.add_paragraph("• Bảng màu, nam châm gắn bảng, thước thẳng.", style='List Bullet')

    add_para(doc, "2. Học sinh:", bold=True)
    doc.add_paragraph("• SGK Toán 5, vở bài tập, bút, thước kẻ.", style='List Bullet')
    doc.add_paragraph("• Đã học bài Diện tích hình thoi (bài trước).", style='List Bullet')
    doc.add_paragraph()

    # CÁC HOẠT ĐỘNG DẠY HỌC
    add_heading(doc, "PHẦN IV: CÁC HOẠT ĐỘNG DẠY HỌC", 1, (0x1a, 0x5f, 0x7a))
    
    add_para(doc, "⏰ Phân bổ thời gian: Khởi động (5') → Khám phá (10') → Hình thành kiến thức (15') → Luyện tập (12') → Vận dụng & Tổng kết (3')", italic=True)
    doc.add_paragraph()

    # Activity table
    act_table = doc.add_table(rows=1, cols=5)
    act_table.style = 'Table Grid'
    headers = ["Hoạt động", "Thời gian", "Nội dung", "Phương pháp / PTDH", "Đánh giá"]
    header_row = act_table.rows[0]
    for i, h_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h_text
        set_cell_bg(cell, "1A5F7A")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    activities = [
        ("HĐ 1:\nKhởi động", "5 phút",
         "GV chiếu hình ảnh mảnh ruộng hình thang ở vùng đồng bằng sông Cửu Long. Hỏi: 'Để tính diện tích mảnh đất này, chúng ta cần biết điều gì?' → Kết nối với bài học.",
         "PP: Đặt câu hỏi gợi mở\nPTDH: Máy chiếu, hình ảnh thực tế\nKỹ thuật: Brainstorm",
         "GV quan sát sự hào hứng, phản hồi nhanh của HS"),
        ("HĐ 2:\nKhám phá hình thang", "10 phút",
         "- GV phát mô hình hình thang cho các nhóm (3-4 HS/nhóm)\n- HS đo và ghi lại: 2 cạnh song song (a, b), chiều cao (h)\n- Đại diện nhóm chia sẻ kết quả\n- GV chuẩn hóa khái niệm và ký hiệu",
         "PP: Học qua trải nghiệm, thảo luận nhóm\nPTDH: Mô hình hình thang, thước\nKỹ thuật: Think-Pair-Share",
         "Phiếu ghi chép của nhóm;\nGV đặt câu hỏi kiểm tra hiểu biết"),
        ("HĐ 3:\nHình thành công thức", "15 phút",
         "- GV hướng dẫn HS cắt ghép hình thang thành hình chữ nhật (bằng mô hình)\n- HS tự phát hiện: S_hcn = chiều dài × chiều rộng\n- Từ đó suy ra: S_thang = (a+b)/2 × h\n- GV viết công thức lên bảng, giải thích từng ký hiệu\n- 2-3 HS nhắc lại công thức",
         "PP: Phát hiện, giảng giải\nPTDH: Mô hình cắt ghép, bảng, phấn màu\nKỹ thuật: Scaffolding",
         "Vấn đáp: GV hỏi từng HS;\nKiểm tra vở ghi công thức"),
        ("HĐ 4:\nLuyện tập", "12 phút",
         "Bài 1: Tính S hình thang có a=10cm, b=6cm, h=4cm\nBài 2: Hình thang có đáy lớn gấp đôi đáy bé, h=8cm, S=120cm². Tìm a, b.\nBài 3 (nâng cao): Bài toán có lời văn về mảnh vườn hình thang\n- HS làm bài cá nhân, sau đó chữa bảng",
         "PP: Thực hành, luyện tập\nPTDH: Vở bài tập, bảng\nKỹ thuật: Chữa bài theo cặp",
         "Chấm nhanh 3-5 vở;\nSửa lỗi sai phổ biến"),
        ("HĐ 5:\nVận dụng & Tổng kết", "3 phút",
         "- GV hỏi: 'Hôm nay em học được gì?'\n- HS đọc lại công thức\n- GV giao BTVN: Bài 1,2 trang 87 SGK\n- Xem trước bài: Diện tích các hình đã học",
         "PP: Tổng kết, củng cố\nPTDH: SGK",
         "HS tự đánh giá;\nGV nhận xét chung lớp"),
    ]
    
    row_colors = ["FDFEFE", "EBF5FB", "FDFEFE", "EBF5FB", "FDFEFE"]
    for i, (act, time, content, method, assess) in enumerate(activities):
        row = act_table.add_row()
        data = [act, time, content, method, assess]
        for j, text in enumerate(data):
            row.cells[j].text = text
            set_cell_bg(row.cells[j], row_colors[i])
            if j == 0:
                for para in row.cells[j].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            else:
                for para in row.cells[j].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    doc.add_paragraph()

    # DỰ KIẾN TÌNH HUỐNG
    add_heading(doc, "PHẦN V: DỰ KIẾN TÌNH HUỐNG SƯ PHẠM", 1, (0x1a, 0x5f, 0x7a))
    
    sit_table = doc.add_table(rows=1, cols=3)
    sit_table.style = 'Table Grid'
    for i, h_text in enumerate(["Tình huống dự kiến", "Phân tích", "Cách xử lý"]):
        cell = sit_table.rows[0].cells[i]
        cell.text = h_text
        set_cell_bg(cell, "1A5F7A")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
    
    situations = [
        ("HS nhầm lẫn chiều cao với cạnh bên hình thang nghiêng",
         "Đây là lỗi phổ biến do HS chưa phân biệt được chiều cao vuông góc với đáy",
         "GV vẽ thêm đường cao h vuông góc, dùng ký hiệu góc vuông, so sánh trực quan với cạnh bên"),
        ("HS không nhớ công thức, tính sai",
         "HS chưa hiểu nguồn gốc công thức, chỉ học thuộc lòng",
         "Cho HS quay lại hoạt động cắt ghép hình thang → hình chữ nhật để tự suy ra công thức"),
        ("Bài toán nâng cao HS không biết bắt đầu từ đâu",
         "HS chưa có kỹ năng phân tích đề bài",
         "GV hướng dẫn: Đọc kỹ đề → Tóm tắt → Vẽ hình → Áp dụng công thức từng bước"),
    ]
    for sit, anal, sol in situations:
        row = sit_table.add_row()
        for j, text in enumerate([sit, anal, sol]):
            row.cells[j].text = text
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # ĐÁNH GIÁ
    add_heading(doc, "PHẦN VI: KẾ HOẠCH ĐÁNH GIÁ", 1, (0x1a, 0x5f, 0x7a))
    
    add_para(doc, "1. Đánh giá quá trình (Formative Assessment):", bold=True)
    doc.add_paragraph("• Quan sát thái độ tham gia hoạt động nhóm.", style='List Bullet')
    doc.add_paragraph("• Đặt câu hỏi vấn đáp trong suốt giờ dạy.", style='List Bullet')
    doc.add_paragraph("• Kiểm tra phiếu học tập nhóm sau hoạt động khám phá.", style='List Bullet')
    
    add_para(doc, "2. Đánh giá kết quả (Summative Assessment):", bold=True)
    doc.add_paragraph("• Chấm bài luyện tập tại lớp (3 bài tập).", style='List Bullet')
    doc.add_paragraph("• Đánh giá bài tập về nhà (bài 1, 2 trang 87).", style='List Bullet')
    
    add_para(doc, "3. Tiêu chí đánh giá kết quả học tập:", bold=True)
    eval_table = doc.add_table(rows=1, cols=3)
    eval_table.style = 'Table Grid'
    for i, h in enumerate(["Mức độ", "Biểu hiện", "Điểm"]):
        set_cell_bg(eval_table.rows[0].cells[i], "2E86C1")
        eval_table.rows[0].cells[i].text = h
        for para in eval_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
    
    eval_data = [
        ("Hoàn thành tốt", "Tính đúng cả 3 bài, trình bày rõ ràng, có thể giải thích được cách làm", "9-10"),
        ("Hoàn thành", "Tính đúng 2/3 bài, trình bày chấp nhận được", "7-8"),
        ("Cần cố gắng", "Chỉ làm đúng bài 1, còn nhầm lẫn trong bài 2, 3", "5-6"),
        ("Chưa hoàn thành", "Không làm được hoặc làm sai cả 3 bài", "Dưới 5"),
    ]
    eval_colors = ["D5F5E3", "D6EAF8", "FDEBD0", "FADBD8"]
    for k, (level, desc, score) in enumerate(eval_data):
        row = eval_table.add_row()
        row.cells[0].text = level
        row.cells[1].text = desc
        row.cells[2].text = score
        set_cell_bg(row.cells[0], eval_colors[k])
        for para in row.cells[2].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    # RÚT KINH NGHIỆM
    add_heading(doc, "PHẦN VII: RÚT KINH NGHIỆM SAU TIẾT DẠY", 1, (0x1a, 0x5f, 0x7a))
    add_para(doc, "(Dành cho giáo viên điền sau khi đã dạy xong tiết học)", italic=True)
    for _ in range(5):
        doc.add_paragraph("_" * 80)
    doc.add_paragraph()

    # TÀI LIỆU THAM KHẢO
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1, (0x1a, 0x5f, 0x7a))
    refs = [
        "1. Bộ Giáo dục và Đào tạo (2024). Sách giáo khoa Toán 5 - Bộ Cánh Diều. NXB Giáo dục Việt Nam.",
        "2. Bộ Giáo dục và Đào tạo (2018). Chương trình Giáo dục phổ thông môn Toán.",
        "3. Nguyễn Bá Kim (2017). Phương pháp dạy học môn Toán. NXB Đại học Sư phạm.",
        "4. Phạm Đình Thực (2019). Thiết kế bài dạy Toán tiểu học. NXB Giáo dục Việt Nam.",
        "5. Bloom, B. S. (1956). Taxonomy of Educational Objectives. Longmans.",
    ]
    for ref in refs:
        add_para(doc, ref, size=10)

    # Signature
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig_para.add_run("TP.HCM, ngày 03 tháng 06 năm 2026\nSinh viên thực hiện\n\n\n\nNguyễn Thị An")
    sig_run.font.size = Pt(11)
    sig_run.bold = True

    filename = os.path.join(OUTPUT_DIR, "BaiLam_01_NguyenThiAn_Xuat_Sac.docx")
    doc.save(filename)
    print(f"✅ Đã tạo: {filename}")


# ============================================================
# BÀI LÀM 2: Khá - Trần Văn Bảo (7.5 điểm)
# ============================================================
def create_assignment_good():
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("TRƯỜNG ĐẠI HỌC SƯ PHẠM TP. HỒ CHÍ MINH - KHOA GDTH\n")
    run.bold = True
    run.font.size = Pt(12)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_para.add_run("GIÁO ÁN DẠY TOÁN LỚP 5")
    t.bold = True
    t.font.size = Pt(14)

    doc.add_paragraph()
    add_para(doc, "Họ tên: Trần Văn Bảo", bold=True)
    add_para(doc, "MSSV: 21CDTH045 | Lớp: GDTH-K21B")
    add_para(doc, "Môn: Kĩ năng dạy học | GV: TS. Trần Văn Bình")
    doc.add_paragraph()

    add_heading(doc, "I. THÔNG TIN BÀI HỌC", 1)
    add_para(doc, "Tên bài: Diện tích hình thang")
    add_para(doc, "Lớp: 5A - Trường Tiểu học Đinh Tiên Hoàng")
    add_para(doc, "Thời gian: 45 phút")
    add_para(doc, "SGK: Toán 5 - Cánh Diều")
    doc.add_paragraph()

    add_heading(doc, "II. MỤC TIÊU", 1)
    add_para(doc, "1. Kiến thức:", bold=True)
    doc.add_paragraph("- Biết công thức tính diện tích hình thang S = (a+b) x h : 2")
    doc.add_paragraph("- Nhận biết các yếu tố của hình thang: đáy lớn, đáy bé, chiều cao")
    
    add_para(doc, "2. Kỹ năng:", bold=True)
    doc.add_paragraph("- Tính được diện tích hình thang qua các bài tập")
    doc.add_paragraph("- Áp dụng vào bài toán có lời văn đơn giản")
    
    add_para(doc, "3. Thái độ:", bold=True)
    doc.add_paragraph("- Học sinh hứng thú với môn học, cẩn thận khi tính toán")
    doc.add_paragraph()

    add_heading(doc, "III. CHUẨN BỊ", 1)
    add_para(doc, "Giáo viên: Giáo án, bảng phụ, thước kẻ, phấn màu", bold=False)
    add_para(doc, "Học sinh: SGK Toán 5, vở bài tập, bút chì")
    doc.add_paragraph()

    add_heading(doc, "IV. TIẾN TRÌNH DẠY HỌC", 1)
    
    proc_table = doc.add_table(rows=1, cols=4)
    proc_table.style = 'Table Grid'
    for i, h in enumerate(["Hoạt động", "Thời gian", "Nội dung & PP", "Ghi chú"]):
        proc_table.rows[0].cells[i].text = h
        set_cell_bg(proc_table.rows[0].cells[i], "2C3E50")
        for para in proc_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)

    activities = [
        ("1. Khởi động", "5'", "GV ôn lại hình thang là hình có 2 cạnh song song. Hỏi HS về đặc điểm. Kết nối vào bài mới.", "Vấn đáp"),
        ("2. Bài mới", "15'", "GV vẽ hình thang lên bảng, ghi a, b, h. Giải thích ý nghĩa từng ký hiệu. Viết công thức: S = (a + b) × h ÷ 2. GV làm mẫu 1 ví dụ với a=12, b=8, h=5.", "Giảng giải, trực quan"),
        ("3. Thực hành", "20'", "Bài 1 (SGK/87): Tính diện tích 3 hình thang. HS làm cá nhân, GV chữa bảng.\nBài 2: Bài toán lời văn - tính diện tích mảnh vườn.", "Thực hành, luyện tập"),
        ("4. Củng cố", "5'", "GV hỏi lại công thức. HS đọc to công thức 3 lần. Giao BTVN trang 87.", "Tổng kết"),
    ]
    
    for act, time, content, method in activities:
        row = proc_table.add_row()
        row.cells[0].text = act
        row.cells[1].text = time
        row.cells[2].text = content
        row.cells[3].text = method
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "V. ĐÁNH GIÁ", 1)
    add_para(doc, "Đánh giá bằng quan sát và chấm bài tập tại lớp. HS làm đúng cả 2 bài đạt yêu cầu.")
    doc.add_paragraph()

    add_heading(doc, "VI. TÀI LIỆU THAM KHẢO", 1)
    add_para(doc, "1. SGK Toán 5 - Cánh Diều (2024)")
    add_para(doc, "2. Vở bài tập Toán 5")

    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.add_run("TP.HCM, ngày 03/06/2026\nSinh viên\n\n\nTrần Văn Bảo").font.size = Pt(11)

    filename = os.path.join(OUTPUT_DIR, "BaiLam_02_TranVanBao_Kha.docx")
    doc.save(filename)
    print(f"✅ Đã tạo: {filename}")


# ============================================================
# BÀI LÀM 3: Trung bình - Lê Thị Cẩm (5.5 điểm)
# ============================================================
def create_assignment_average():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GIÁO ÁN TOÁN LỚP 5\nDIỆN TÍCH HÌNH THANG")
    r.bold = True
    r.font.size = Pt(14)
    
    doc.add_paragraph()
    add_para(doc, "Sinh viên: Lê Thị Cẩm - MSSV: 21CDTH078")
    add_para(doc, "Lớp GDTH-K21C - Kĩ năng dạy học")
    doc.add_paragraph()

    add_heading(doc, "I. MỤC TIÊU", 1)
    add_para(doc, "- Học sinh biết tính diện tích hình thang")
    add_para(doc, "- Nhớ được công thức")
    add_para(doc, "- Làm được bài tập trong SGK")
    doc.add_paragraph()

    add_heading(doc, "II. ĐỒ DÙNG", 1)
    add_para(doc, "SGK, vở, bảng, phấn")
    doc.add_paragraph()

    add_heading(doc, "III. CÁC BƯỚC LÊN LỚP", 1)
    
    add_para(doc, "1. Kiểm tra bài cũ (3 phút):", bold=True)
    add_para(doc, "Hỏi HS: Nêu đặc điểm của hình thang. Gọi 1 HS lên bảng.")
    
    add_para(doc, "2. Bài mới (30 phút):", bold=True)
    add_para(doc, "- GV viết đề bài lên bảng: Diện tích hình thang")
    add_para(doc, "- GV vẽ hình thang lên bảng, ghi a=10, b=6, h=4")
    add_para(doc, "- GV nói: Công thức tính diện tích hình thang là S = (a+b) x h : 2")
    add_para(doc, "- GV tính mẫu: S = (10+6) x 4 : 2 = 32 cm²")
    add_para(doc, "- Cho HS làm bài 1 trong SGK")
    add_para(doc, "- Gọi HS lên bảng chữa bài")
    add_para(doc, "- GV nhận xét")
    
    add_para(doc, "3. Củng cố - Dặn dò (2 phút):", bold=True)
    add_para(doc, "Nhắc lại công thức. Về nhà làm bài tập còn lại.")
    doc.add_paragraph()

    add_heading(doc, "IV. RÚT KINH NGHIỆM", 1)
    add_para(doc, "(Bỏ trống)")
    doc.add_paragraph()
    
    # Thiếu phần tài liệu tham khảo và ký tên

    filename = os.path.join(OUTPUT_DIR, "BaiLam_03_LeThiCam_TrungBinh.docx")
    doc.save(filename)
    print(f"✅ Đã tạo: {filename}")


# ============================================================
# BÀI LÀM 4: Yếu - Phạm Minh Đức (3.5 điểm)
# ============================================================
def create_assignment_weak():
    doc = Document()

    title = doc.add_paragraph("GIÁO ÁN")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Môn: Toán lớp 5")
    doc.add_paragraph("Bài: Hình thang")
    doc.add_paragraph("Giáo viên: Phạm Minh Đức")
    doc.add_paragraph()
    
    doc.add_paragraph("Mục tiêu:")
    doc.add_paragraph("Học sinh biết hình thang")
    doc.add_paragraph("Tính được diện tích")
    doc.add_paragraph()

    doc.add_paragraph("Chuẩn bị:")
    doc.add_paragraph("SGK và bảng")
    doc.add_paragraph()
    
    doc.add_paragraph("Tiến hành:")
    doc.add_paragraph("1. Ôn bài cũ: hỏi bài cũ")
    doc.add_paragraph("2. Bài mới: GV dạy về hình thang. Công thức: S = (a+b)*h/2. Ví dụ: S = (5+3)*4/2 = 16")
    doc.add_paragraph("3. Luyện tập: cho HS làm bài trong sách")
    doc.add_paragraph("4. Dặn dò: học bài")
    doc.add_paragraph()
    doc.add_paragraph("(Không có phần đánh giá, không có tài liệu tham khảo)")

    filename = os.path.join(OUTPUT_DIR, "BaiLam_04_PhamMinhDuc_Yeu.docx")
    doc.save(filename)
    print(f"✅ Đã tạo: {filename}")


# ============================================================
# BÀI LÀM 5: Giỏi - Hoàng Phương Linh (8.5 điểm)
# ============================================================
def create_assignment_very_good():
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = header_para.add_run("TRƯỜNG ĐẠI HỌC SƯ PHẠM TP. HỒ CHÍ MINH\n")
    r1.bold = True; r1.font.size = Pt(13)
    r2 = header_para.add_run("BÀI DỰ THI THIẾT KẾ GIÁO ÁN MÔN TOÁN LỚP 5\n")
    r2.bold = True; r2.font.size = Pt(11)

    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_para.add_run("DIỆN TÍCH HÌNH THANG\n(Tích hợp phương pháp Dạy học dựa trên vấn đề - PBL)")
    t.bold = True; t.font.size = Pt(15)
    t.font.color.rgb = RGBColor(0x15, 0x6B, 0x39)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run("Sinh viên: Hoàng Phương Linh | MSSV: 21CDTH022 | Lớp: GDTH-K21A\n").bold = True
    info.add_run("Môn: Kĩ năng dạy học | GV hướng dẫn: TS. Trần Văn Bình").font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "I. THÔNG TIN CHUNG", 1, (0x15, 0x6B, 0x39))
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_rows = [
        ("Tên bài:", "Diện tích hình thang - SGK Toán 5, Bài 31"),
        ("Lớp - Trường:", "5B - Trường Tiểu học Nguyễn Trãi, Q.1, TP.HCM"),
        ("Thời lượng:", "45 phút"),
        ("Phương pháp chủ đạo:", "Dạy học dựa trên vấn đề (PBL) + Học hợp tác"),
        ("Chuẩn kiến thức:", "Theo CTGDPT 2018 - Cấp tiểu học"),
    ]
    for key, val in info_rows:
        row = info_table.add_row()
        row.cells[0].text = key
        row.cells[1].text = val
        set_cell_bg(row.cells[0], "E9F7EF")
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "II. MỤC TIÊU THEO BLOOM", 1, (0x15, 0x6B, 0x39))
    add_para(doc, "Mục tiêu được xây dựng theo thang nhận thức Bloom (2001):", italic=True)
    
    bloom_table = doc.add_table(rows=1, cols=3)
    bloom_table.style = 'Table Grid'
    for i, h in enumerate(["Cấp độ Bloom", "Mục tiêu", "Động từ đo lường"]):
        bloom_table.rows[0].cells[i].text = h
        set_cell_bg(bloom_table.rows[0].cells[i], "1E8449")
        for para in bloom_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
    
    bloom_data = [
        ("Nhớ (Remember)", "Ghi nhớ công thức S = (a+b)×h÷2 và tên các yếu tố", "Kể tên, liệt kê, nhắc lại"),
        ("Hiểu (Understand)", "Giải thích ý nghĩa từng ký hiệu trong công thức", "Giải thích, mô tả"),
        ("Vận dụng (Apply)", "Tính diện tích hình thang khi biết a, b, h", "Tính toán, giải bài"),
        ("Phân tích (Analyze)", "Phân biệt chiều cao với cạnh bên, xác định đúng các đáy", "Phân biệt, xác định"),
    ]
    bloom_colors = ["D5F5E3", "D6EAF8", "FEF9E7", "FDEDEC"]
    for k, (level, obj, verb) in enumerate(bloom_data):
        row = bloom_table.add_row()
        row.cells[0].text = level; row.cells[1].text = obj; row.cells[2].text = verb
        set_cell_bg(row.cells[0], bloom_colors[k])
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "III. CHUẨN BỊ", 1, (0x15, 0x6B, 0x39))
    add_para(doc, "Giáo viên:", bold=True)
    items_gv = ["Giáo án chi tiết, slide PowerPoint (15 slide)", 
                "Mô hình hình thang bằng bìa màu (5 bộ, mỗi bộ 4 hình)",
                "Phiếu học tập A3 in sẵn bài toán thực tế (tình huống mảnh vườn)",
                "Bảng từ, nam châm, thước đo góc, thước thẳng",
                "Bảng rubric đánh giá nhóm (in cho mỗi HS 1 tờ)"]
    for item in items_gv:
        doc.add_paragraph(f"• {item}")
    add_para(doc, "Học sinh:", bold=True)
    for item in ["SGK Toán 5, vở ô ly, bút màu, thước kẻ", "Đã học hình thang, công thức hình chữ nhật"]:
        doc.add_paragraph(f"• {item}")
    doc.add_paragraph()

    add_heading(doc, "IV. TÌNH HUỐNG XUẤT PHÁT (Bối cảnh PBL)", 1, (0x15, 0x6B, 0x39))
    context_para = doc.add_paragraph()
    context_para.add_run(
        "🌾 Tình huống thực tế: Gia đình bạn Nam có một mảnh vườn hình thang ở nông thôn. "
        "Bố Nam muốn trồng rau trên mảnh đất đó và cần biết diện tích để mua hạt giống cho đủ. "
        "Mảnh đất có đáy lớn 24m, đáy bé 16m, chiều cao 12m. "
        "Em hãy giúp bạn Nam tính diện tích mảnh vườn!"
    ).italic = True
    doc.add_paragraph()

    add_heading(doc, "V. TIẾN TRÌNH DẠY HỌC CHI TIẾT", 1, (0x15, 0x6B, 0x39))
    
    act_table = doc.add_table(rows=1, cols=4)
    act_table.style = 'Table Grid'
    for i, h in enumerate(["Hoạt động", "TG", "Mô tả chi tiết", "Đánh giá"]):
        act_table.rows[0].cells[i].text = h
        set_cell_bg(act_table.rows[0].cells[i], "1E8449")
        for para in act_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
    
    activities = [
        ("Khởi động\n& Đặt vấn đề", "7'",
         "GV kể câu chuyện bạn Nam và mảnh vườn. Chiếu hình ảnh ruộng hình thang.\nHỏi: 'Tại sao cần biết diện tích?' → HS nêu ý kiến.\nGV giới thiệu: 'Hôm nay chúng ta sẽ giúp bạn Nam!'",
         "Quan sát sự tham gia; Đặt câu hỏi mở"),
        ("Khám phá\nhình thang", "8'",
         "Chia nhóm 5 người, phát mô hình hình thang.\nNhóm: Đo a, b, h; Ghi vào phiếu; So sánh với hình chữ nhật.\nĐại diện nhóm trình bày kết quả đo.",
         "Kiểm tra phiếu nhóm;\nVấn đáp trực tiếp"),
        ("Xây dựng\ncông thức", "10'",
         "HS cắt hình thang, ghép lại thành hình chữ nhật → Thảo luận.\nGV hỏi: 'S hình thang = ? so với S hình chữ nhật?'\nDẫn dắt: S = (a+b)×h÷2\nGV ghi công thức, HS ghi vào vở.",
         "Vấn đáp; Kiểm tra vở ghi"),
        ("Luyện tập\nphân hóa", "15'",
         "Mức 1: Tính S với số cho sẵn (3 bài)\nMức 2: Tìm chiều cao khi biết S và hai đáy\nMức 3 (nâng cao): Giải bài toán bạn Nam + so sánh 2 mảnh vườn\nHS chọn mức, làm cá nhân, chữa theo cặp.",
         "GV chấm nhanh;\nKhen ngợi kịp thời"),
        ("Kết luận\n& Ứng dụng", "5'",
         "Nhóm trình bày đáp án bài toán bạn Nam.\nCả lớp nhận xét, bổ sung.\nGV tổng kết: ứng dụng trong đo đất, kiến trúc.\nGiao BTVN có liên hệ thực tế.",
         "Tự đánh giá theo rubric;\nGV nhận xét tổng thể"),
    ]
    
    for act, time, desc, assess in activities:
        row = act_table.add_row()
        row.cells[0].text = act; row.cells[1].text = time
        row.cells[2].text = desc; row.cells[3].text = assess
        for j, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if j == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()

    add_heading(doc, "VI. PHÂN HÓA DẠY HỌC", 1, (0x15, 0x6B, 0x39))
    diff_table = doc.add_table(rows=1, cols=3)
    diff_table.style = 'Table Grid'
    for i, h in enumerate(["HS gặp khó khăn", "HS đạt chuẩn", "HS tiến bộ/năng khiếu"]):
        diff_table.rows[0].cells[i].text = h
        colors = ["E74C3C", "2980B9", "27AE60"]
        set_cell_bg(diff_table.rows[0].cells[i], colors[i][1:])
        for para in diff_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
    
    diff_row = diff_table.add_row()
    diff_row.cells[0].text = "• Cung cấp gợi ý từng bước\n• Dùng bảng công thức\n• Bài toán số nhỏ, đơn giản"
    diff_row.cells[1].text = "• Làm đầy đủ 3 bài mức 1\n• Thử bài mức 2\n• Tham gia thảo luận nhóm"
    diff_row.cells[2].text = "• Làm thêm bài mức 3\n• Tự đặt đề bài mới\n• Hỗ trợ bạn trong nhóm"
    for cell in diff_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "VII. ĐÁNH GIÁ", 1, (0x15, 0x6B, 0x39))
    add_para(doc, "Sử dụng đánh giá đa dạng:", bold=True)
    doc.add_paragraph("• Đánh giá quá trình: Quan sát thái độ, tham gia thảo luận, phiếu nhóm.")
    doc.add_paragraph("• Đánh giá sản phẩm: Bài làm 3 mức độ phân hóa.")
    doc.add_paragraph("• Tự đánh giá: HS dùng rubric cá nhân đánh giá bài làm của mình.")
    doc.add_paragraph("• Đánh giá đồng đẳng: Chữa bài theo cặp, nhóm nhận xét nhóm.")
    doc.add_paragraph()

    add_heading(doc, "VIII. TÀI LIỆU THAM KHẢO", 1, (0x15, 0x6B, 0x39))
    refs = [
        "1. Bộ GD&ĐT (2024). SGK Toán 5 - Cánh Diều. NXB Giáo dục.",
        "2. Bộ GD&ĐT (2018). Chương trình GDPT môn Toán.",
        "3. Nguyễn Bá Kim (2017). Phương pháp dạy học Toán. NXB ĐHSP.",
        "4. Barrows, H.S. (1986). A taxonomy of problem-based learning methods. Medical Education.",
        "5. Bloom, B.S. et al. (2001). A taxonomy for learning, teaching, and assessing.",
    ]
    for ref in refs:
        add_para(doc, ref, size=10)

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run("TP.HCM, ngày 03 tháng 06 năm 2026\n\n\nHoàng Phương Linh").bold = True

    filename = os.path.join(OUTPUT_DIR, "BaiLam_05_HoangPhuongLinh_Gioi.docx")
    doc.save(filename)
    print(f"✅ Đã tạo: {filename}")


if __name__ == "__main__":
    print("🎓 Đang tạo các bài làm mẫu...\n")
    create_assignment_excellent()
    create_assignment_very_good()
    create_assignment_good()
    create_assignment_average()
    create_assignment_weak()
    print(f"\n✅ Hoàn thành! Đã tạo 5 bài làm mẫu trong thư mục: {OUTPUT_DIR}")
