"""PDF and DOCX text extraction."""
import io
import re

try:
    import PyPDF2
    _HAS_PYPDF2 = True
except ImportError:
    _HAS_PYPDF2 = False

try:
    import docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


def clean_extracted_pdf_text(text: str) -> str:
    """
    Heuristic to fix PDF extraction layout bugs (e.g. PyPDF2 extracting Vietnamese word-by-word on separate lines).
    If the number of lines is high relative to the number of words, merge single newlines.
    """
    if not text:
        return ""
    words = text.split()
    if not words:
        return text
    lines = text.splitlines()
    
    # If more than 60% of words are on their own lines, it's likely a word-by-word layout bug
    if len(lines) > 0.6 * len(words):
        # Standardize newlines
        cleaned = text.replace("\r\n", "\n")
        
        # Heuristic to distinguish word/line wraps from paragraph breaks.
        # Word transitions in buggy PDFs usually have 1 or 2 newlines (e.g. \n \n).
        # Real paragraph breaks have 3 or more newlines (e.g. \n \n \n).
        def replace_newline_sequence(match):
            seq = match.group(0)
            newline_count = seq.count('\n')
            if newline_count >= 3:
                return "\n\n"
            else:
                return " "
        
        cleaned = re.sub(r'\n[\s\n]*', replace_newline_sequence, cleaned)
        # Remove consecutive spaces
        cleaned = re.sub(r' +', ' ', cleaned)
        return cleaned.strip()
    return text


def extract_text_from_pdf(file_path: str) -> str:
    if not _HAS_PYPDF2:
        return ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = [page.extract_text() or "" for page in reader.pages]
        raw_text = "\n".join(pages).strip()
        return clean_extracted_pdf_text(raw_text)
    except Exception:
        return ""


def extract_text_from_pdf_bytes(data: bytes) -> str:
    if not _HAS_PYPDF2:
        return ""
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        raw_text = "\n".join(pages).strip()
        return clean_extracted_pdf_text(raw_text)
    except Exception:
        return ""


def extract_text_from_docx_bytes(data: bytes) -> str:
    if not _HAS_DOCX:
        return ""
    try:
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        doc = docx.Document(io.BytesIO(data))
        
        full_text_blocks = []
        for element in doc.element.body:
            tag_name = element.tag.split('}')[-1]
            if tag_name == 'p':
                p = Paragraph(element, doc)
                p_text = p.text.strip()
                if p_text:
                    full_text_blocks.append(p_text)
            elif tag_name == 'tbl':
                table = Table(element, doc)
                table_lines = []
                for row in table.rows:
                    row_text = []
                    prev_cell = None
                    for cell in row.cells:
                        if cell == prev_cell:
                            continue
                        cell_txt = cell.text.strip()
                        if cell_txt:
                            row_text.append(cell_txt)
                        prev_cell = cell
                    if row_text:
                        table_lines.append(" | ".join(row_text))
                if table_lines:
                    full_text_blocks.append("\n".join(table_lines))
                    
        if not full_text_blocks:
            # Fallback to paragraph and table separate if ordered extraction returned nothing
            paras_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    prev_cell = None
                    for cell in row.cells:
                        if cell == prev_cell:
                            continue
                        cell_txt = cell.text.strip()
                        if cell_txt:
                            row_text.append(cell_txt)
                        prev_cell = cell
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            tbls_text = "\n".join(tables_text)
            if paras_text:
                full_text_blocks.append(paras_text)
            if tbls_text:
                full_text_blocks.append(tbls_text)
                
        return "\n\n".join(full_text_blocks).strip()
    except Exception:
        # Simplest fallback
        try:
            doc = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
        except Exception:
            return ""


def extract_text_from_file(filename: str, data: bytes) -> str:
    """Auto-detect file type and extract text."""
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf_bytes(data)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx_bytes(data)
    elif name_lower.endswith(".txt"):
        try:
            return data.decode("utf-8").strip()
        except Exception:
            try:
                return data.decode("latin-1").strip()
            except Exception:
                return ""
    return ""
